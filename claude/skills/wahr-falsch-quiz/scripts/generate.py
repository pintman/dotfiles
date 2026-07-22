#!/usr/bin/env python3
"""Erzeugt ein Wahr-Falsch-Quiz (.docx) im Microsoft-Forms-Import-Format aus einer JSON-Konfiguration.

Aufruf:
    python generate.py --config config.json [--output pfad.docx] [--variante schueler|lehrer|beide]

Siehe example_config.json für das erwartete Schema.
"""

import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Pt

INTRO = (
    "Beurteile die Korrektheit der folgenden Aussagen. "
    "Für jede korrekte Antwort gibt es einen Punkt, für jede falsche Antwort gibt es einen Minuspunkt. "
    "Durch den Abzug können nicht weniger als 0 Punkte in dieser Aufgabe erzielt werden. "
    "Jede nicht oder nicht eindeutig gekennzeichnete Aussage ergibt 0 Punkte. "
    "Rate nicht."
)


def load_config(path):
    with open(path, encoding='utf-8') as f:
        config = json.load(f)

    required = ('aussagen', 'output')
    missing = [key for key in required if key not in config]
    if missing:
        raise ValueError(f'Config fehlt Pflichtfelder: {", ".join(missing)}')

    for i, aussage in enumerate(config['aussagen'], start=1):
        if 'text' not in aussage or 'antwort' not in aussage:
            raise ValueError(f'Aussage {i} fehlt "text" oder "antwort"')
        if aussage['antwort'] not in ('wahr', 'falsch'):
            raise ValueError(f'Aussage {i}: "antwort" muss "wahr" oder "falsch" sein')

    return config


def build_document(aussagen, mit_loesung):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_paragraph(INTRO)
    doc.add_paragraph('')

    for i, aussage in enumerate(aussagen, start=1):
        text = aussage['text']
        if mit_loesung:
            praefix = '[w] ' if aussage['antwort'] == 'wahr' else '[f] '
            text = praefix + text
        doc.add_paragraph(f'{i}. {text}')
        doc.add_paragraph('a) Wahr')
        doc.add_paragraph('b) Falsch')
        doc.add_paragraph('')

    return doc


def loesung_pfad(output_path):
    root, ext = os.path.splitext(output_path)
    return f'{root}_LOESUNG{ext}'


def save(doc, path):
    out_dir = os.path.dirname(path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(path)
    print(f'Gespeichert: {path}')


def main():
    parser = argparse.ArgumentParser(description='Erzeugt ein Wahr-Falsch-Quiz (.docx) aus einer JSON-Konfiguration.')
    parser.add_argument('--config', required=True, help='Pfad zur JSON-Konfigurationsdatei')
    parser.add_argument('--output', help='Ausgabepfad (.docx); überschreibt "output" aus der Config')
    parser.add_argument(
        '--variante', choices=['schueler', 'lehrer', 'beide'], default='lehrer',
        help='schueler = ohne Lösungspräfix (Forms-Import), lehrer = mit [w]/[f]-Präfix (Standard), '
             'beide = beide Dateien erzeugen',
    )
    args = parser.parse_args()

    config = load_config(args.config)
    aussagen = config['aussagen']
    output_path = args.output or config['output']

    if args.variante == 'schueler':
        save(build_document(aussagen, mit_loesung=False), output_path)
    elif args.variante == 'lehrer':
        save(build_document(aussagen, mit_loesung=True), output_path)
    else:
        save(build_document(aussagen, mit_loesung=False), output_path)
        save(build_document(aussagen, mit_loesung=True), loesung_pfad(output_path))


if __name__ == '__main__':
    sys.exit(main())
