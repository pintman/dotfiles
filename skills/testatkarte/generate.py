#!/usr/bin/env python3
"""Erzeugt eine Testatkarte (.docx) aus einer JSON-Konfiguration.

Aufruf:
    python generate.py --config config.json [--output pfad.docx]

Siehe example_config.json für das erwartete Schema.
"""

import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Twips, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Konstanten ──────────────────────────────────────────────
PAGE_W = 11906
PAGE_H = 16838
MARGIN = 1080  # ~1.9 cm – etwas enger für mehr Platz
CONTENT_W = PAGE_W - 2 * MARGIN  # 9746

BK = '000000'  # Schwarz
WH = 'FFFFFF'  # Weiß
GR = 'EEEEEE'  # Hellgrau für Code

THIN_SZ = 2
THICK_SZ = 8

COL_NR = 480
COL_TESTAT = 1800


# ── Hilfsfunktionen ─────────────────────────────────────────
def set_cell_borders(cell, top=THIN_SZ, bottom=THIN_SZ, left=THIN_SZ, right=THIN_SZ):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, sz in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), BK)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for edge, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def set_cell_width(cell, width):
    cell.width = Twips(width)


def style_run(run, text, size=18, bold=False, italic=False, font='Arial'):
    run.text = text
    run.font.name = font
    run.font.size = Pt(size / 2)
    run.font.bold = bold
    run.font.italic = italic


def header_cell(cell, text, width, align=WD_ALIGN_PARAGRAPH.CENTER):
    set_cell_width(cell, width)
    set_cell_borders(cell, top=THICK_SZ, bottom=THICK_SZ)
    set_cell_shading(cell, WH)
    set_cell_margins(cell, top=80, bottom=80)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    style_run(para.add_run(), text, size=18, bold=True)


def data_cell(cell, text, width, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    set_cell_width(cell, width)
    set_cell_borders(cell)
    set_cell_shading(cell, WH)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    style_run(para.add_run(), text, size=17, bold=bold)


def add_heading(doc, text, level=1):
    sizes = {1: 28, 2: 24, 3: 20}
    spacing_before = {1: 240, 2: 180, 3: 120}
    spacing_after = {1: 120, 2: 80, 3: 60}
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Twips(spacing_before[level])
    para.paragraph_format.space_after = Twips(spacing_after[level])
    style_run(para.add_run(), text, size=sizes[level], bold=True)
    if level == 1:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), BK)
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return para


def add_p(doc, text, center=False, before=0, after=100, code=False, size=None, bold=False, italic=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Twips(before)
    para.paragraph_format.space_after = Twips(after)
    if code:
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), GR)
        pPr.append(shd)
    font = 'Courier New' if code else 'Arial'
    default_size = 16 if code else 18
    style_run(para.add_run(), text, size=size or default_size, bold=bold, italic=italic, font=font)
    return para


def add_code_line(doc, text):
    return add_p(doc, text, code=True, after=0)


def add_prompt_text(doc, text):
    return add_p(doc, text, italic=True, size=17, after=60)


def add_page_break(doc):
    doc.add_page_break()


def add_class_diagram(doc, title, attributes, methods, width=4000):
    table = doc.add_table(rows=3, cols=1)
    table.autofit = False

    head = table.rows[0].cells[0]
    set_cell_width(head, width)
    set_cell_borders(head, top=THICK_SZ, bottom=THICK_SZ, left=THICK_SZ, right=THICK_SZ)
    set_cell_margins(head, top=60, bottom=60)
    head.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(head.paragraphs[0].add_run(), title, size=20, bold=True)

    attrs_cell = table.rows[1].cells[0]
    set_cell_width(attrs_cell, width)
    set_cell_borders(attrs_cell, top=THIN_SZ, bottom=THIN_SZ, left=THICK_SZ, right=THICK_SZ)
    set_cell_margins(attrs_cell, top=60, bottom=60)
    for line in attributes:
        para = attrs_cell.add_paragraph() if attrs_cell.paragraphs[0].runs else attrs_cell.paragraphs[0]
        style_run(para.add_run(), line, size=18, font='Courier New')

    methods_cell = table.rows[2].cells[0]
    set_cell_width(methods_cell, width)
    set_cell_borders(methods_cell, top=THIN_SZ, bottom=THIN_SZ, left=THICK_SZ, right=THICK_SZ)
    set_cell_margins(methods_cell, top=60, bottom=60)
    for line in methods:
        para = methods_cell.add_paragraph() if methods_cell.paragraphs[0].runs else methods_cell.paragraphs[0]
        style_run(para.add_run(), line, size=18, font='Courier New')

    return table


# ── Konfiguration laden ──────────────────────────────────────
def load_config(path):
    with open(path, encoding='utf-8') as f:
        config = json.load(f)

    required = ('thema', 'klasse', 'einleitung', 'meilensteine', 'output')
    missing = [key for key in required if key not in config]
    if missing:
        raise ValueError(f'Config fehlt Pflichtfelder: {", ".join(missing)}')

    config.setdefault('zusatzinfos', [])
    return config


# ── Seite 1: Namenszeile, Titel, Meilensteintabelle ─────────
def build_page1(doc, config):
    name_table = doc.add_table(rows=1, cols=2)
    name_table.autofit = False
    name_cells = name_table.rows[0].cells
    set_cell_width(name_cells[0], 1600)
    set_cell_borders(name_cells[0], top=THICK_SZ, bottom=THICK_SZ, left=THICK_SZ, right=THICK_SZ)
    set_cell_margins(name_cells[0], top=80, bottom=80)
    style_run(name_cells[0].paragraphs[0].add_run(), 'Name:', size=20, bold=True)

    set_cell_width(name_cells[1], CONTENT_W - 1600)
    set_cell_borders(name_cells[1], top=THICK_SZ, bottom=THICK_SZ, left=THICK_SZ, right=THICK_SZ)
    set_cell_margins(name_cells[1], top=80, bottom=80)
    style_run(name_cells[1].paragraphs[0].add_run(), '', size=20)

    add_p(doc, '')

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Twips(160)
    style_run(title_para.add_run(), f'Testatkarte – {config["thema"]} – {config["klasse"]}', size=36, bold=True)

    intro_para = doc.add_paragraph()
    intro_para.paragraph_format.space_after = Twips(180)
    style_run(intro_para.add_run(), config['einleitung'], size=17)

    meilensteine = config['meilensteine']
    col_text = CONTENT_W - COL_NR - COL_TESTAT

    m_table = doc.add_table(rows=1 + len(meilensteine), cols=3)
    m_table.autofit = False

    header_row = m_table.rows[0].cells
    header_cell(header_row[0], 'Nr.', COL_NR, WD_ALIGN_PARAGRAPH.CENTER)
    header_cell(header_row[1], 'Meilenstein', col_text, WD_ALIGN_PARAGRAPH.LEFT)
    header_cell(header_row[2], 'Testat, Datum', COL_TESTAT, WD_ALIGN_PARAGRAPH.CENTER)

    for i, m in enumerate(meilensteine, start=1):
        row = m_table.rows[i].cells
        data_cell(row[0], str(m['nr']), COL_NR, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        data_cell(row[1], m['text'], col_text)
        data_cell(row[2], '', COL_TESTAT)


# ── Seiten 2–4: Zusatzinformationen aus Blocks ───────────────
def build_zusatzinfos(doc, blocks):
    for block in blocks:
        btype = block['type']
        if btype == 'pagebreak':
            add_page_break(doc)
        elif btype == 'heading':
            add_heading(doc, block['text'], level=block.get('level', 2))
        elif btype == 'paragraph':
            add_p(
                doc, block['text'],
                center=block.get('center', False),
                before=block.get('before', 0),
                after=block.get('after', 100),
                size=block.get('size'),
                bold=block.get('bold', False),
                italic=block.get('italic', False),
            )
        elif btype == 'code':
            for line in block['lines']:
                add_code_line(doc, line)
        elif btype == 'prompt':
            for line in block['lines']:
                add_prompt_text(doc, line)
        elif btype == 'classdiagram':
            add_class_diagram(doc, block['title'], block.get('attributes', []), block.get('methods', []))
        elif btype == 'spacer':
            add_p(doc, '')
        else:
            raise ValueError(f'Unbekannter Block-Typ: {btype!r}')


def build_document(config):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Twips(PAGE_W)
    section.page_height = Twips(PAGE_H)
    section.top_margin = Twips(MARGIN)
    section.bottom_margin = Twips(MARGIN)
    section.left_margin = Twips(MARGIN)
    section.right_margin = Twips(MARGIN)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    build_page1(doc, config)

    if config['zusatzinfos']:
        add_page_break(doc)
        build_zusatzinfos(doc, config['zusatzinfos'])

    return doc


def main():
    parser = argparse.ArgumentParser(description='Erzeugt eine Testatkarte (.docx) aus einer JSON-Konfiguration.')
    parser.add_argument('--config', required=True, help='Pfad zur JSON-Konfigurationsdatei')
    parser.add_argument('--output', help='Ausgabepfad (.docx); überschreibt "output" aus der Config')
    args = parser.parse_args()

    config = load_config(args.config)
    output_path = args.output or config['output']

    doc = build_document(config)

    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    print(f'Fertig: {output_path}')


if __name__ == '__main__':
    sys.exit(main())
